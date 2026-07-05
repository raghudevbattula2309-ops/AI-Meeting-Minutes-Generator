"""
Word Export Service
"""

from io import BytesIO
from docx import Document


def export_minutes_to_word(minutes):
    """
    Generate a Word document and return it as bytes.
    """

    document = Document()

    document.add_heading(
        "AI Engineering Meeting Minutes",
        level=1
    )

    document.add_paragraph()

    for line in minutes.splitlines():

        line = line.strip()

        if not line:
            continue

        if line.startswith("# "):

            document.add_heading(
                line.replace("# ", ""),
                level=1
            )

        elif line.startswith("## "):

            document.add_heading(
                line.replace("## ", ""),
                level=2
            )

        else:

            document.add_paragraph(line)

    buffer = BytesIO()

    document.save(buffer)

    buffer.seek(0)

    return buffer